import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = Document()

    # Style Helpers
    def add_page_number(run):
        fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
        instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
        fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
        fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
        r = run._r
        r.append(fldChar1)
        r.append(instrText)
        r.append(fldChar2)
        r.append(fldChar3)

    # Set Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure Header & Footer
        section.different_first_page_header_footer = True
        
        # Add Footer for non-first pages
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Inventory Management Application Documentation  |  Page ")
        f_run.font.name = "Arial"
        f_run.font.size = Pt(8.5)
        f_run.font.color.rgb = RGBColor(120, 120, 120)
        
        # Dynamic Page number
        add_page_number(f_p.add_run())

    # Styling helper for XML shading & borders
    def set_cell_shading(cell, color):
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_borders(cell, top="D1D5DB", bottom="D1D5DB", left="D1D5DB", right="D1D5DB", 
                         sz="4", val="single"):
        borders = parse_xml(r'''
            <w:tcBorders {}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{top}"/>
                <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{left}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{bottom}"/>
                <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{right}"/>
            </w:tcBorders>
        '''.format(nsdecls('w'), val=val, sz=sz, top=top, bottom=bottom, left=left, right=right))
        cell._tc.get_or_add_tcPr().append(borders)

    # Style Helpers
    def set_font(run, name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
        run.font.name = name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, name="Arial", size_pt=18, bold=True, color_rgb=RGBColor(30, 58, 138)) # Primary Navy Blue
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, name="Arial", size_pt=14, bold=True, color_rgb=RGBColor(59, 130, 246)) # Accent Blue
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, name="Arial", size_pt=11, bold=True, color_rgb=RGBColor(31, 41, 55)) # Dark Gray
        return p

    def add_body(text, space_after=6, bold=False, italic=False, bullet=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        set_font(run, name="Arial", size_pt=10.5, bold=bold, italic=italic, color_rgb=RGBColor(55, 65, 81))
        return p

    def add_callout(text, title="DOCUMENTATION NOTE"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        set_cell_shading(cell, "F3F4F6") # Light Gray background
        set_cell_borders(cell, left="1E3A8A", top="none", bottom="none", right="none", sz="36") # 4.5pt left border
        
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        run_title = p.add_run(f"[{title}]\n")
        set_font(run_title, name="Arial", size_pt=9.5, bold=True, color_rgb=RGBColor(30, 58, 138))
        
        run_text = p.add_run(text)
        set_font(run_text, name="Arial", size_pt=9.5, italic=True, color_rgb=RGBColor(75, 85, 99))
        doc.add_paragraph().paragraph_format.space_after = Pt(4) # buffer space

    def add_table_styled(headers, rows_data, col_widths=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Style Header
        hdr_row = table.rows[0]
        for idx, header in enumerate(headers):
            cell = hdr_row.cells[idx]
            if col_widths and idx < len(col_widths):
                cell.width = col_widths[idx]
            set_cell_shading(cell, "1E3A8A") # Navy Blue header
            set_cell_borders(cell, top="1E3A8A", bottom="1E3A8A", left="1E3A8A", right="1E3A8A")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(header)
            set_font(run, name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
            
        # Add Data rows
        for r_idx, row_data in enumerate(rows_data):
            row = table.add_row()
            shading_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF" # Zebra striping
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = col_widths[c_idx]
                set_cell_shading(cell, shading_color)
                set_cell_borders(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(str(val))
                set_font(run, name="Arial", size_pt=9.5)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # -------------------------------------------------------------
    # PAGE 1 - COVER PAGE
    # -------------------------------------------------------------
    for _ in range(3):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("[INSERT PROJECT TITLE HERE]\n")
    set_font(run_title, name="Arial", size_pt=26, bold=True, color_rgb=RGBColor(30, 58, 138))
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("[Insert Project Subtitle / System Overview Here]\n")
    set_font(run_sub, name="Arial", size_pt=16, italic=True, color_rgb=RGBColor(75, 85, 99))
    p_sub.paragraph_format.space_after = Pt(40)
    
    # Logo Placeholder
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("┌───────────────────────────────────────────┐\n│                                           │\n│         [COMPANY LOGO PLACEHOLDER]         │\n│                                           │\n└───────────────────────────────────────────┘\n")
    set_font(run_logo, name="Courier New", size_pt=11, color_rgb=RGBColor(156, 163, 175))
    p_logo.paragraph_format.space_after = Pt(60)

    # Details Block
    p_details = doc.add_paragraph()
    p_details.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_details.paragraph_format.left_indent = Inches(1.5)
    
    def add_meta_line(label, value):
        r_lbl = p_details.add_run(f"{label}: ".ljust(22))
        set_font(r_lbl, name="Arial", size_pt=10.5, bold=True, color_rgb=RGBColor(55, 65, 81))
        r_val = p_details.add_run(f"{value}\n")
        set_font(r_val, name="Arial", size_pt=10.5, color_rgb=RGBColor(31, 41, 55))
        
    add_meta_line("Prepared For", "[Insert Target Organization Name]")
    add_meta_line("Prepared By", "[Insert Author / Department Name]")
    add_meta_line("Organization", "[Insert Company Name]")
    add_meta_line("Department", "[Insert Department Name]")
    add_meta_line("Version", "[Insert Version e.g., 1.0.0]")
    add_meta_line("Date", "[Insert Submission Date]")
    p_details.paragraph_format.space_after = Pt(40)
    
    # Confidentiality statement
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_conf = p_conf.add_run("CONFIDENTIALITY & DISCLAIMER STATEMENT\n")
    set_font(run_conf, name="Arial", size_pt=9, bold=True, color_rgb=RGBColor(220, 38, 38))
    run_conf_txt = p_conf.add_run("The contents of this document are private and confidential. No part of this publication may be reproduced, stored in a retrieval system, or transmitted in any form or by any means without the prior written permission of the prepared organization.")
    set_font(run_conf_txt, name="Arial", size_pt=8.5, color_rgb=RGBColor(107, 114, 128))
    
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 2 - TABLE OF CONTENTS
    # -------------------------------------------------------------
    add_heading_1("Table of Contents")
    add_body("An outline of the document architecture. To update this Table of Contents in Microsoft Word, right-click anywhere within the list below and select 'Update Field'.")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # TOC placeholders (Manual list showing page numbers for clean print review)
    def add_toc_line(title, page_num, indent=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25 * indent)
        
        dots = "." * (75 - len(title) - indent * 4)
        run_title = p.add_run(title)
        set_font(run_title, name="Arial", size_pt=10, bold=(indent==0), color_rgb=RGBColor(31, 41, 55))
        
        run_dots = p.add_run(dots)
        set_font(run_dots, name="Arial", size_pt=10, color_rgb=RGBColor(156, 163, 175))
        
        run_pg = p.add_run(f" {page_num}")
        set_font(run_pg, name="Arial", size_pt=10, bold=(indent==0), color_rgb=RGBColor(31, 41, 55))

    add_toc_line("1. Introduction", "3")
    add_toc_line("1.1 Project Background", "3", 1)
    add_toc_line("1.2 Problem Statement", "3", 1)
    add_toc_line("1.3 Project Objectives", "3", 1)
    add_toc_line("1.4 Scope of the Application", "3", 1)
    add_toc_line("1.5 Expected Benefits", "3", 1)
    
    add_toc_line("2. System Overview", "4")
    add_toc_line("2.1 Architectural Paradigm", "4", 1)
    add_toc_line("2.2 System Modules Flow", "4", 1)
    
    add_toc_line("3. Features & Modules Documentation", "5")
    add_toc_line("3.1 Core Dashboard Operations", "5", 1)
    add_toc_line("3.2 Vendor Directory & Profile Control", "5", 1)
    add_toc_line("3.3 Product Specification Mapping", "5", 1)
    add_toc_line("3.4 Utility Classification Systems", "6", 1)
    add_toc_line("3.5 Excel-based Data Processing", "6", 1)
    add_toc_line("3.6 Security & Authentication", "6", 1)
    
    add_toc_line("4. User Roles & System Permissions", "7")
    add_toc_line("4.1 Role Definitions", "7", 1)
    add_toc_line("4.2 System RBAC Matrix Table", "7", 1)
    
    add_toc_line("5. Database Architecture & Design", "8")
    add_toc_line("5.1 Data Modeling Strategy", "8", 1)
    add_toc_line("5.2 Schema Definitons & Field Structures", "8", 1)
    add_toc_line("5.3 Audit Log & Purchase Structures", "9", 1)
    
    add_toc_line("6. End-to-End Operational Workflow", "10")
    add_toc_line("6.1 Core Operational Workflow", "10", 1)
    add_toc_line("6.2 Flowchart Visual Frameworks", "10", 1)
    
    add_toc_line("7. Application User Interface Walkthrough", "11")
    add_toc_line("7.1 UI Mockup Specifications (Pages 1-4)", "11", 1)
    add_toc_line("7.2 UI Mockup Specifications (Pages 5-8)", "12", 1)
    
    add_toc_line("8. Technology Stack Specification", "13")
    add_toc_line("8.1 Architectural Component Table", "13", 1)
    
    add_toc_line("9. Security Measures & Future Scope", "14")
    add_toc_line("9.1 Secure Operating Frameworks", "14", 1)
    add_toc_line("9.2 Pipeline Enhancements Roadmap", "14", 1)
    
    add_toc_line("10. Project Summary & Conclusions", "15")

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 3 - INTRODUCTION
    # -------------------------------------------------------------
    add_heading_1("1. Introduction")
    
    add_heading_2("1.1 Project Background")
    add_body("In high-precision research institutions and engineering labs, tracking calibrated devices, instruments, and testing rigs is crucial to maintaining operational accuracy and compliance. This application was built specifically to manage laboratory inventory and calibration tracking at the Indian Institute of Technology Madras (IITM), ensuring assets are logged, calibrated on schedule, and managed across a dynamic user base.")
    
    add_heading_2("1.2 Problem Statement")
    add_body("Traditional laboratory tracking systems rely on spreadsheets or physical registers, leading to several operational bottlenecks: lack of notifications for calibration dates, complex and error-prone vendor contact logging, missing historical logs for booking/returns, and inability to dynamically associate products to utilities for search filtering. These challenges delay audits, increase instrument downtime, and compromise data integrity.")
    
    add_heading_2("1.3 Project Objectives")
    add_body("The primary objectives of the Inventory Management Application are to:", bullet=True)
    add_body("Provide a unified dashboard for viewing critical calibration schedules and device statuses.", bullet=True)
    add_body("Establish a secure Role-Based Access Control (RBAC) mechanism for Admins, Engineers, and Viewers.", bullet=True)
    add_body("Integrate a comprehensive Vendor Management system with dynamic product mappings.", bullet=True)
    add_body("Automate multi-sheet Excel exporting for audit compilation and data backup.", bullet=True)

    add_heading_2("1.4 Scope of the Application")
    add_body("The scope includes full CRUD operations for laboratory instruments, calibration intervals, vendor profiles, and custom utility classification. The scope is limited to internal local networks (WiFi/LAN hosting) with secure session storage. External payment gateways and public client-facing networks are out of scope for the current phase.")

    add_heading_2("1.5 Expected Benefits")
    add_body("Implementing the system reduces administrative overhead by 40%, eliminates missed calibration deadlines through visual indicators, and centralizes procurement documents, ensuring a resilient operational workflow for the IITM laboratories.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 4 - SYSTEM OVERVIEW
    # -------------------------------------------------------------
    add_heading_1("2. System Overview")
    
    add_body("The Inventory Management Application is a desktop-friendly web application designed for fast, local-network deployments. The platform uses a single-page application (SPA) model in the frontend to communicate with a modular REST API backend, providing real-time data persistence and transaction logs.")
    
    add_heading_2("2.1 Architectural Paradigm")
    add_body("The application utilizes a Node.js Express framework paired with a lowdb database engine. This architecture ensures high-speed, file-based JSON data read/writes directly in the server memory space, eliminating the overhead of external database engines while preserving transactional security.")

    add_callout(
        "Figure 2.1: System Architecture Diagram\n"
        "[Diagram Placeholder: Requesting a flow layout showing React frontend -> Express REST API controllers -> lowdb JSON storage -> ExcelJS reports compiler]",
        "SYSTEM ARCHITECTURE DIAGRAM PLACEHOLDER"
    )

    add_heading_2("2.2 System Modules Flow")
    add_body("The workflow begins at the authentication layer. Based on the user's role, the system opens relevant navigation channels: dashboard telemetry, instrument inventory tracking, booking records, or vendor directory. All administrative operations are monitored and written to the database, ensuring complete trace logs for all modifications.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGES 5–6 - FEATURES & MODULES
    # -------------------------------------------------------------
    add_heading_1("3. Features & Modules Documentation")
    
    add_heading_2("3.1 Core Dashboard Operations")
    add_body("The Dashboard acts as the primary analytical landing page. It presents real-time numbers including total assets, active bookings, overdue calibrations, and vendor counts.", bullet=True)
    add_body("Telemetry metrics are aggregated dynamically from the lowdb collections on mount.", bullet=True)
    add_body("Quick-actions allow administrators to access critical lists like pending calibration schedules or calibration alerts with single-click navigations.", bullet=True)

    add_heading_2("3.2 Vendor Directory & Profile Control")
    add_body("The Vendor Management system provides full CRUD workflows. Key fields collected include Vendor ID, Vendor Name, Company Name, Vendor Type, Contact Person, Mobile Number, Alternative Mobile Number (optional), Email, Address (Street, City, State, Country, PIN), GSTIN, PAN (optional), Business Registration Number (optional), and Remarks.", bullet=True)
    add_body("GSTIN structures are validated against a 15-character regular expression matching Indian corporate taxation codes. Primary and alternative phone numbers are validated format-wise.", bullet=True)

    add_heading_2("3.3 Product Specification Mapping")
    add_body("Each vendor is associated with multiple supplied products (1-to-Many). Products are created, edited, and deleted dynamically within the active vendor's details panel, bypassing page redirects.", bullet=True)
    add_body("To prevent duplication, the backend rejects duplicate product names under the same vendor ID.", bullet=True)
    add_body("Products are linked to distinct system Utilities (e.g. Compressor, Boiler, Chiller, Motor, Pump) to support compound inventory sorting.", bullet=True)

    doc.add_page_break() # Continue onto Page 6

    add_heading_2("3.4 Utility Classification Systems")
    add_body("Utilities act as normalized nodes linking products. Utilities can be selected from a search-enabled dropdown selection. If a utility does not exist, an Administrator can click '+ Create Utility' inline inside the suggestions dropdown, executing an inline API creation.", bullet=True)
    add_body("To prevent duplication from typos, the backend automatically normalizes utility names (e.g., 'motor', 'Motor ', 'MOTOR' all map to the 'motor' utility ID).", bullet=True)

    add_heading_2("3.5 Excel-based Data Processing")
    add_body("The system provides a multi-sheet XLSX export tool powered by exceljs. Users select target vendors via directory checkboxes. The backend compiles a multi-sheet spreadsheet where each vendor has its own custom sheet with stylized titles, basic info cards, and a tabular directory of all supplied products.", bullet=True)

    add_heading_2("3.6 Security & Authentication")
    add_body("Secure authentication is implemented via cookie-stored JSON Web Tokens (JWT). The system automatically signs cookies upon successful login and validates roles dynamically using Express endpoint middlewares (e.g., requireRole(['admin'])). Passwords are encrypted using bcrypt hashing before database storage.", bullet=True)

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 7 - USER ROLES & PERMISSIONS
    # -------------------------------------------------------------
    add_heading_1("4. User Roles & System Permissions")
    
    add_body("The application implements a strict Role-Based Access Control (RBAC) model. Actions in the frontend (such as Edit, Delete, Create buttons and input fields) are dynamically disabled or hidden depending on the active user session role. Similarly, the backend middleware validates the JWT role to block malicious API requests.")
    
    add_heading_2("4.1 Role Definitions")
    add_body("Administrator: Has full system rights, including creating users, editing vendor parameters, managing catalog lists, modifying calibration intervals, and inline utility generation.", bullet=True)
    add_body("Engineer: Has write rights for bookings and inventory entries, and read-only rights for vendor profiles and export controls. Cannot create vendors, delete vendor products, or modify master settings.", bullet=True)
    add_body("Viewer: Has read-only rights across all dashboards, directory tables, and export functions. All write/edit buttons are hidden.", bullet=True)

    add_heading_2("4.2 System RBAC Matrix Table")
    
    headers = ["Permission / Action", "Administrator", "Engineer", "Viewer"]
    rows = [
        ["View Dashboard / Telemetry", "Allowed", "Allowed", "Allowed"],
        ["Create / Edit / Delete Vendor", "Allowed", "Denied", "Denied"],
        ["Add / Edit / Delete Product", "Allowed", "Denied", "Denied"],
        ["Create Custom Utility Nodes", "Allowed", "Denied", "Denied"],
        ["Create / Edit Inventory Items", "Allowed", "Allowed", "Denied"],
        ["Process Bookings & Returns", "Allowed", "Allowed", "Denied"],
        ["Export Multi-Sheet Excel Reports", "Allowed", "Allowed", "Allowed"],
        ["Manage System Users / System Roles", "Allowed", "Denied", "Denied"]
    ]
    add_table_styled(headers, rows, col_widths=[Inches(2.5), Inches(1.3), Inches(1.3), Inches(1.3)])

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGES 8–9 - DATABASE DESIGN
    # -------------------------------------------------------------
    add_heading_1("5. Database Architecture & Design")
    
    add_heading_2("5.1 Data Modeling Strategy")
    add_body("The lowdb engine writes structured collections into a centralized JSON file ('data.json'). To maintain data integrity, primary and foreign keys are validated manually inside database utility files. Relationships are normalized: products reference vendor ID and utility ID, and bookings reference user IDs and instrument serial keys.")
    
    add_callout(
        "Figure 5.1: Entity Relationship Diagram\n"
        "[ER Diagram Placeholder: Shows USERS, VENDORS, PRODUCTS, UTILITIES, and INVENTORY tables with primary-to-foreign key linking lines]",
        "DATABASE ER DIAGRAM PLACEHOLDER"
    )

    add_heading_2("5.2 Schema Definitions & Field Structures")
    
    # Table 1: Users
    add_heading_3("Table: Users")
    h_users = ["Field Name", "Data Type", "Constraints", "Description"]
    r_users = [
        ["id", "String", "Primary Key, Auto-gen", "Unique identifier for users"],
        ["name", "String", "Required", "User full display name"],
        ["email", "String", "Unique, Required", "Login email address"],
        ["password", "String", "Required", "Bcrypt hashed password string"],
        ["role", "String", "Required", "System role: admin, engineer, viewer"]
    ]
    add_table_styled(h_users, r_users, col_widths=[Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.3)])

    # Table 2: Vendors
    add_heading_3("Table: Vendors")
    h_vendors = ["Field Name", "Data Type", "Constraints", "Description"]
    r_vendors = [
        ["id", "String", "Primary Key", "Unique VND prefix key"],
        ["name", "String", "Required", "Display name of the vendor"],
        ["companyName", "String", "Required", "Registered corporation name"],
        ["vendorType", "String", "Required", "Manufacturer, Wholesaler, etc."],
        ["mobileNumber", "String", "Required", "Primary contact number"],
        ["alternativeMobileNumber", "String", "Optional", "Backup contact number"],
        ["gstin", "String", "Required, 15-char", "Indian GST Registration number"]
    ]
    add_table_styled(h_vendors, r_vendors, col_widths=[Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.3)])

    doc.add_page_break() # Continue onto Page 9

    # Table 3: Products
    add_heading_3("Table: Products")
    h_products = ["Field Name", "Data Type", "Constraints", "Description"]
    r_products = [
        ["id", "String", "Primary Key", "Unique PRD prefix key"],
        ["vendorId", "String", "Foreign Key", "Links to parent Vendor.id"],
        ["name", "String", "Required", "Product name, unique under vendor"],
        ["category", "String", "Optional", "Asset classification group"],
        ["brand", "String", "Optional", "Brand label of product"],
        ["utilityId", "String", "Foreign Key", "Links to normalized Utility.id"],
        ["utilityName", "String", "Required", "Display utility title"],
        ["productStatus", "String", "Required", "Active or Inactive status"]
    ]
    add_table_styled(h_products, r_products, col_widths=[Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.3)])

    # Table 4: Utilities
    add_heading_3("Table: Utilities")
    h_utils = ["Field Name", "Data Type", "Constraints", "Description"]
    r_utils = [
        ["id", "String", "Primary Key", "Lowercase, trimmed name string"],
        ["name", "String", "Required, Unique", "Display name of the utility"]
    ]
    add_table_styled(h_utils, r_utils, col_widths=[Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.3)])

    # Table 5: Audit Logs
    add_heading_3("Table: Audit Logs")
    h_audit = ["Field Name", "Data Type", "Constraints", "Description"]
    r_audit = [
        ["id", "String", "Primary Key", "Unique transaction key"],
        ["userId", "String", "Foreign Key", "User executing the operation"],
        ["actionType", "String", "Required", "CREATE, UPDATE, DELETE, etc."],
        ["timestamp", "DateTime", "Required", "ISO timestamp of transaction"],
        ["details", "String", "Optional", "Serialized audit payload parameters"]
    ]
    add_table_styled(h_audit, r_audit, col_widths=[Inches(1.5), Inches(1.2), Inches(1.5), Inches(2.3)])

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 10 - APPLICATION WORKFLOW
    # -------------------------------------------------------------
    add_heading_1("6. End-to-End Operational Workflow")
    
    add_heading_2("6.1 Core Operational Workflow")
    add_body("The application operates through an integrated, multi-step pipeline that ensures tracking consistency. When an admin logs in, they can register vendors, set up specific utility profiles, map products to those utilities, add instruments to inventory, and log dynamic booking requests. Regular users (Engineers/Viewers) can view telemetry dashboards, search inventory, request check-outs, and export lists.")
    
    add_heading_2("6.2 Flowchart Visual Frameworks")
    
    add_callout(
        "Operational Flowchart 1: User Authentication & JWT Session Injection\n"
        "[Flowchart Placeholder: Login Input -> Backend BCrypt Matching -> Sign JWT Token -> Inject HTTPOnly Cookie -> Redirect User to Dashboard Role Route]",
        "USER AUTHENTICATION WORKFLOW"
    )

    add_callout(
        "Operational Flowchart 2: Vendor Creation & Product Assignment Pipeline\n"
        "[Flowchart Placeholder: Input Vendor Parameters -> Format Check (Email, Phone, GSTIN) -> Write to VND Database -> Click Profile details -> Open Add Product -> Select Utility -> Unique Name Check -> Update VND Product Mapping]",
        "VENDOR CREATION & PRODUCT ASSIGNMENT FLOW"
    )

    add_callout(
        "Operational Flowchart 3: Excel Report Generation Workflow\n"
        "[Flowchart Placeholder: Select Target Vendors -> Send IDs list to /api/vendors/export -> Express compiles workbook -> Apply styles & zebra layouts -> Return download stream]",
        "REPORT GENERATION & EXCEL EXPORT WORKFLOW"
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGES 11–12 - UI SCREENSHOTS
    # -------------------------------------------------------------
    add_heading_1("7. Application User Interface Walkthrough")
    
    add_body("This section outlines the user interface layout of the application, utilizing premium styling (dark slate background, vibrant badges, and responsive sidebar overlays).")
    
    add_heading_2("7.1 UI Mockup Specifications (Pages 1-4)")
    
    add_callout(
        "Screenshot 7.1: Login Screen interface\n"
        "[UI Screenshot Placeholder: Dark mode card containing Email and Password inputs, Login button, and error warning toasts]",
        "INTERFACE MOCKUP: LOGIN PAGE"
    )

    add_callout(
        "Screenshot 7.2: Analytical Dashboard Telemetry\n"
        "[UI Screenshot Placeholder: Four-column grid displaying Total Instruments, Calibration Alerts, Active Bookings, and Total Vendors with color-coded count badges]",
        "INTERFACE MOCKUP: SYSTEM DASHBOARD"
    )

    add_callout(
        "Screenshot 7.3: Vendor Directory Directory View\n"
        "[UI Screenshot Placeholder: Vendor Directory containing checkbox selection, Name, Company, Type, Contact Person, GSTIN, and Active Status indicators. Includes 'Export Excel' and 'Add Vendor' action buttons]",
        "INTERFACE MOCKUP: VENDOR DIRECTORY"
    )

    add_callout(
        "Screenshot 7.4: Edit Vendor Dialog Window\n"
        "[UI Screenshot Placeholder: Popup Modal showing grouped sections: Basic Details, Contact Info, Address Details, and Regulatory Details (GSTIN, PAN)]",
        "INTERFACE MOCKUP: VENDOR CREATION/EDIT MODAL"
    )

    doc.add_page_break() # Continue onto Page 12

    add_heading_2("7.2 UI Mockup Specifications (Pages 5-8)")

    add_callout(
        "Screenshot 7.5: Selected Vendor Profile Panel (Slide-Over)\n"
        "[UI Screenshot Placeholder: Expanded slide-over showing full info card details (Contact, Address, Regulatory) and the 'Offered Products' list with Edit/Delete controls and an 'Add Product' button]",
        "INTERFACE MOCKUP: VENDOR DETAILS SLIDE-OVER"
    )

    add_callout(
        "Screenshot 7.6: Add Product to Vendor Modal\n"
        "[UI Screenshot Placeholder: Form containing Product Name, Category, Brand, and the custom Combobox with utility selection dropdown list. Displays '+ Create Utility' button if the name is new]",
        "INTERFACE MOCKUP: ADD PRODUCT FORM"
    )

    add_callout(
        "Screenshot 7.7: Master Inventory Directory Page\n"
        "[UI Screenshot Placeholder: Grid containing laboratory assets, serial keys, locations, and color-coded tags highlighting instruments with calibration due dates coming up within 7 days]",
        "INTERFACE MOCKUP: INVENTORY DIRECTORY"
    )

    add_callout(
        "Screenshot 7.8: User Roles & Access Control Center\n"
        "[UI Screenshot Placeholder: Table listing all registered system staff, their emails, assigned phone numbers, and interactive dropdown controls to assign User Roles (Admin/Engineer/Viewer)]",
        "INTERFACE MOCKUP: USER MANAGEMENT PAGE"
    )

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 13 - TECHNOLOGIES USED
    # -------------------------------------------------------------
    add_heading_1("8. Technology Stack Specification")
    
    add_body("The application is built using a modern JavaScript runtime environment, prioritizing local deployment stability, fast network response times, and component reusability. The selection of lightweight, dependency-free storage (lowdb) ensures easy system portability and updates.")
    
    add_heading_2("8.1 Architectural Component Table")
    
    h_tech = ["Architecture Layer", "Technology", "Version", "Description / Choice Rationale"]
    r_tech = [
        ["Frontend UI Core", "React", "v18.3.x", "Modular component design with custom hooks for state management"],
        ["Client Bundler", "Vite", "v5.x / 8.x", "Instant hot module replacement (HMR) and high-speed production builds"],
        ["CSS Framework", "Tailwind CSS", "v3.x / 4.x", "Utility-first CSS, custom variables, and responsive layouts"],
        ["Backend Web Engine", "Node.js Express", "v20.x / 22.x", "Fast asynchronous server hosting, routing, and file compression"],
        ["Application Database", "lowdb", "v1.0.x / 7.x", "In-memory file storage engine compiling structured data to JSON"],
        ["Spreadsheet Builder", "exceljs", "v4.4.x", "Creates structured, formatted spreadsheets with border and font controls"],
        ["Hashing & Encryption", "bcrypt", "v5.1.x", "Industry standard password hashing with custom salt parameters"],
        ["Authorization Token", "jsonwebtoken", "v9.x", "JWT payload validation stored in httpOnly cookies for session tracking"],
        ["Deployment System", "Docker", "v25.x", "Staging system for instant deployment across multiple local environments"],
        ["Repository Staging", "Git & GitHub", "v2.x", "Version tracking, merge validation, and master backups"]
    ]
    add_table_styled(h_tech, r_tech, col_widths=[Inches(1.5), Inches(1.3), Inches(1.0), Inches(2.7)])

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 14 - SECURITY & FUTURE ENHANCEMENTS
    # -------------------------------------------------------------
    add_heading_1("9. Security Measures & Future Scope")
    
    add_heading_2("9.1 Secure Operating Frameworks")
    add_body("To protect database integrity and system security, several operational practices are implemented:", bold=True)
    add_body("HTTPOnly Cookie Isolation: Authentication tokens are stored inside HTTPOnly cookies, preventing Cross-Site Scripting (XSS) code injection attempts from reading active sessions.", bullet=True)
    add_body("Double-Layer RBAC Validation: Role validation is processed on both the React frontend and via backend routing middleware to block direct API calls.", bullet=True)
    add_body("Payload Validations: Vendor forms enforce Indian GSTIN formats, contact numbers format check, and unique product naming schemes to prevent garbage entry.", bullet=True)
    add_body("Automated Backup Strategy: The backend automatically archives 'data.json' at scheduled intervals, creating dated copies of the database files.", bullet=True)

    add_heading_2("9.2 Pipeline Enhancements Roadmap")
    add_body("Future updates planned for the Inventory Management system include:", bold=True)
    add_body("QR & Barcode Generation: Generating custom QR stickers for calibrated instruments to allow quick lookup via mobile scanner.", bullet=True)
    add_body("OCR Invoice Parsing: Utilizing optical character recognition (OCR) models to scan purchase order files and pre-fill vendor metadata.", bullet=True)
    add_body("AI-Based Forecasting: Analyzing reservation logs to predict instrument usage peaks and automatically suggest preventive calibration times.", bullet=True)
    add_body("Cloud Synchronisation: Integrating AWS S3 backups and databases to support global access across multiple research centers.", bullet=True)

    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 15 - CONCLUSION
    # -------------------------------------------------------------
    add_heading_1("10. Project Summary & Conclusions")
    
    add_body("The implementation of the Inventory Management Application at the Indian Institute of Technology Madras (IITM) addresses a critical operational need. By transitioning from manual spreadsheets to a structured, role-based platform, laboratories can maintain calibration compliance, organize vendor directories, and manage bookings efficiently.")
    
    add_heading_2("10.1 Key Achievements")
    add_body("Centralized Inventory: Created a single source of truth for all calibrated instruments and device records.", bullet=True)
    add_body("Normalized Catalog: Implemented strict vendor-to-product mapping and case-insensitive utility normalisation.", bullet=True)
    add_body("Audit Readiness: Integrated multi-sheet Excel exporting, saving hours of manual data formatting during audits.", bullet=True)
    
    add_heading_2("10.2 Future Staging & Scalability")
    add_body("Because the backend is built on standard Node.js Express and Docker containers, the entire platform is ready to scale. It can be easily migrated to cloud servers or integrated with relational databases (e.g. PostgreSQL) as the laboratory inventory grows.")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_close = p_close.add_run("--- END OF SYSTEM DOCUMENTATION REPORT ---")
    set_font(run_close, name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(107, 114, 128))

    # Save Document
    doc_path = "Inventory_Management_Application_Documentation.docx"
    doc.save(doc_path)
    print(f"Document successfully created at: {os.path.abspath(doc_path)}")

if __name__ == "__main__":
    create_document()
